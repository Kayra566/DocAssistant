"""ExportDocument → dosya baytları (MD, DOCX, XLSX, PDF)."""

from __future__ import annotations

import io
import os
import re
from xml.sax.saxutils import escape

from app.exports.builder import ExportDocument, Section, Table
from app.models.collab import ExportFormat

MIME_TYPES: dict[ExportFormat, str] = {
    ExportFormat.MD: "text/markdown; charset=utf-8",
    ExportFormat.DOCX: (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ),
    ExportFormat.XLSX: (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    ),
    ExportFormat.PDF: "application/pdf",
}

# Excel hücrelerinde kabul edilmeyen kontrol karakterleri.
_ILLEGAL_XLSX = re.compile(r"[\x00-\x08\x0b-\x0c\x0e-\x1f]")


def _clean(value: str) -> str:
    return _ILLEGAL_XLSX.sub("", str(value))


# ---------- Markdown ----------
def render_markdown(doc: ExportDocument) -> bytes:
    lines: list[str] = [f"# {doc.title}", "", f"_{doc.subtitle}_", ""]
    for section in doc.sections:
        if section.heading:
            lines += [f"## {section.heading}", ""]
        for para in section.paragraphs:
            lines += [para, ""]
        for bullet in section.bullets:
            lines.append(f"- {bullet}")
        if section.bullets:
            lines.append("")
        if section.table:
            lines += _markdown_table(section.table)
    return "\n".join(lines).encode("utf-8")


def _markdown_table(table: Table) -> list[str]:
    header = "| " + " | ".join(table.columns) + " |"
    divider = "| " + " | ".join("---" for _ in table.columns) + " |"
    body = [
        "| " + " | ".join(cell.replace("|", "\\|") for cell in row) + " |"
        for row in table.rows
    ]
    return [header, divider, *body, ""]


# ---------- DOCX ----------
def render_docx(doc: ExportDocument) -> bytes:
    from docx import Document as DocxDocument

    docx = DocxDocument()
    docx.add_heading(doc.title, level=0)
    docx.add_paragraph(doc.subtitle)

    for section in doc.sections:
        if section.heading:
            docx.add_heading(section.heading, level=1)
        for para in section.paragraphs:
            docx.add_paragraph(para)
        for bullet in section.bullets:
            docx.add_paragraph(bullet, style="List Bullet")
        if section.table:
            table = docx.add_table(rows=1, cols=len(section.table.columns))
            table.style = "Table Grid"
            for cell, name in zip(table.rows[0].cells, section.table.columns, strict=True):
                cell.text = name
            for row in section.table.rows:
                cells = table.add_row().cells
                for cell, value in zip(cells, row, strict=True):
                    cell.text = value

    buffer = io.BytesIO()
    docx.save(buffer)
    return buffer.getvalue()


# ---------- XLSX ----------
def render_xlsx(doc: ExportDocument) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    sheet = wb.active
    sheet.title = "Sonuç"
    sheet.append([_clean(doc.title)])
    sheet["A1"].font = Font(bold=True, size=14)
    sheet.append([_clean(doc.subtitle)])
    sheet.append([])

    for section in doc.sections:
        if section.heading:
            sheet.append([_clean(section.heading)])
            sheet.cell(row=sheet.max_row, column=1).font = Font(bold=True)
        for para in section.paragraphs:
            sheet.append([_clean(para)])
        for bullet in section.bullets:
            sheet.append([f"• {_clean(bullet)}"])
        if section.table:
            sheet.append([_clean(c) for c in section.table.columns])
            for col in range(1, len(section.table.columns) + 1):
                sheet.cell(row=sheet.max_row, column=col).font = Font(bold=True)
            for row in section.table.rows:
                sheet.append([_clean(cell) for cell in row])
        sheet.append([])

    sheet.column_dimensions["A"].width = 60
    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


# ---------- PDF ----------
_FONT_NAME = "DocSans"
_FONT_BOLD = "DocSans-Bold"
_fonts_ready = False


def _register_fonts() -> None:
    """reportlab ile gelen Vera fontlarını kaydeder (Türkçe glifleri kapsar)."""
    global _fonts_ready
    if _fonts_ready:
        return
    import reportlab
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    fonts_dir = os.path.join(os.path.dirname(reportlab.__file__), "fonts")
    pdfmetrics.registerFont(TTFont(_FONT_NAME, os.path.join(fonts_dir, "Vera.ttf")))
    pdfmetrics.registerFont(TTFont(_FONT_BOLD, os.path.join(fonts_dir, "VeraBd.ttf")))
    pdfmetrics.registerFontFamily(_FONT_NAME, normal=_FONT_NAME, bold=_FONT_BOLD)
    _fonts_ready = True


def _pdf_styles():
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet

    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "DocTitle", parent=base["Title"], fontName=_FONT_BOLD, fontSize=18
        ),
        "subtitle": ParagraphStyle(
            "DocSubtitle",
            parent=base["Normal"],
            fontName=_FONT_NAME,
            fontSize=9,
            textColor="#666666",
        ),
        "heading": ParagraphStyle(
            "DocHeading",
            parent=base["Heading2"],
            fontName=_FONT_BOLD,
            fontSize=13,
            spaceBefore=12,
        ),
        "body": ParagraphStyle(
            "DocBody", parent=base["Normal"], fontName=_FONT_NAME, fontSize=10, leading=14
        ),
        "bullet": ParagraphStyle(
            "DocBullet",
            parent=base["Normal"],
            fontName=_FONT_NAME,
            fontSize=10,
            leading=14,
            leftIndent=14,
            bulletIndent=4,
        ),
    }


def _pdf_table(table: Table, styles) -> list:
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, TableStyle
    from reportlab.platypus import Table as PdfTable

    header = [Paragraph(f"<b>{escape(c)}</b>", styles["body"]) for c in table.columns]
    body = [
        [Paragraph(escape(cell), styles["body"]) for cell in row] for row in table.rows
    ]
    pdf_table = PdfTable([header, *body], repeatRows=1)
    pdf_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eeeeee")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return [pdf_table]


def _pdf_section(section: Section, styles) -> list:
    from reportlab.platypus import Paragraph, Spacer

    flow: list = []
    if section.heading:
        flow.append(Paragraph(escape(section.heading), styles["heading"]))
    for para in section.paragraphs:
        flow.append(Paragraph(escape(para).replace("\n", "<br/>"), styles["body"]))
        flow.append(Spacer(1, 6))
    for bullet in section.bullets:
        flow.append(Paragraph(escape(bullet), styles["bullet"], bulletText="•"))
    if section.table and section.table.rows:
        flow.append(Spacer(1, 6))
        flow.extend(_pdf_table(section.table, styles))
    flow.append(Spacer(1, 10))
    return flow


def render_pdf(doc: ExportDocument) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    _register_fonts()
    styles = _pdf_styles()
    buffer = io.BytesIO()
    template = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        title=doc.title,
        leftMargin=48,
        rightMargin=48,
        topMargin=48,
        bottomMargin=48,
    )

    flow: list = [
        Paragraph(escape(doc.title), styles["title"]),
        Paragraph(escape(doc.subtitle), styles["subtitle"]),
        Spacer(1, 14),
    ]
    for section in doc.sections:
        flow.extend(_pdf_section(section, styles))

    template.build(flow)
    return buffer.getvalue()


RENDERERS = {
    ExportFormat.MD: render_markdown,
    ExportFormat.DOCX: render_docx,
    ExportFormat.XLSX: render_xlsx,
    ExportFormat.PDF: render_pdf,
}


def render(doc: ExportDocument, fmt: ExportFormat) -> bytes:
    return RENDERERS[fmt](doc)
